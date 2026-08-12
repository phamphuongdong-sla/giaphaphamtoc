import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, text, bold_prefix=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border only
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="C9923A"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
        r_bold.font.name = 'Arial'
        r_bold.font.size = Pt(10.5)
        r_bold.font.color.rgb = RGBColor(140, 20, 20)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Arial'
    r_text.font.size = Pt(10.5)
    r_text.font.color.rgb = RGBColor(50, 50, 50)

def main():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Styles
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    
    run_title = title_p.add_run("📖 HƯỚNG DẪN CÀI ĐẶT VÀ SỬ DỤNG\nỨNG DỤNG GIA PHẢ PHẠM TỘC")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(139, 26, 26) # Dark Burgundy/Red

    # Link Callout
    add_callout(
        doc,
        "https://phamphuongdong-sla.github.io/giaphaphamtoc/\n(Có thể truy cập trực tiếp trên Điện thoại, Máy tính bảng và Máy tính)",
        "🌐 Địa chỉ truy cập Gia Phả: "
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PHẦN 1 ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("📱 PHẦN 1: HƯỚNG DẪN CÀI ĐẶT LÊN MÀN HÌNH CHÍNH (PWA)")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = RGBColor(139, 26, 26)

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(8)
    r = p_intro.add_run(
        "Ứng dụng Gia Phả Phạm Tộc hỗ trợ công nghệ Progressive Web App (PWA), cho phép cài đặt trực tiếp lên thiết bị "
        "mà không cần vào App Store hay Google Play Store. Sau khi cài đặt, ứng dụng sẽ có biểu tượng riêng trên màn hình chính, "
        "mở nhanh chóng và xem được ngay cả khi không có mạng (Offline)."
    )
    r.font.name = 'Arial'
    r.font.size = Pt(11)

    # 1.1 iOS
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("1.1. Trên điện thoại iPhone / iPad (iOS)")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    add_callout(doc, "Bắt buộc phải mở link bằng trình duyệt Safari mặc định của iPhone.", "⚠️ Lưu ý quan trọng: ")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    steps_ios = [
        ("Bước 1: ", "Mở trình duyệt Safari và truy cập đường link gia phả: https://phamphuongdong-sla.github.io/giaphaphamtoc/"),
        ("Bước 2: ", "Bấm vào nút Chia sẻ (biểu tượng hình ô vuông có mũi tên chỉ lên ⬆️) nằm ở thanh công cụ dưới cùng màn hình."),
        ("Bước 3: ", "Vuốt danh sách tùy chọn xuống dưới và chọn dòng \"Thêm vào MH chính\" (Add to Home Screen)."),
        ("Bước 4: ", "Bấm \"Thêm\" (Add) ở góc trên bên phải màn hình."),
        ("🎉 Kết quả: ", "Biểu tượng Gia Phả Phạm Tộc 🏛️ sẽ xuất hiện ngoài màn hình chính của iPhone/iPad như ứng dụng thật.")
    ]
    for b, text in steps_ios:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_b = p.add_run(b)
        r_b.bold = True
        r_b.font.name = 'Arial'
        r_b.font.size = Pt(11)
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(11)

    # 1.2 Android
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("1.2. Trên điện thoại Android (Samsung, Xiaomi, OPPO, Vivo, Realme...)")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    p_c1 = doc.add_paragraph()
    p_c1.paragraph_format.space_after = Pt(4)
    r = p_c1.add_run("Cách 1 (Nhanh nhất):")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)

    steps_android_1 = [
        "Mở link bằng trình duyệt Google Chrome.",
        "Sau 2 - 3 giây, ở dưới cùng màn hình sẽ xuất hiện bảng thông báo: \"Cài Gia Phả lên điện thoại\".",
        "Bấm nút \"Cài đặt\"."
    ]
    for text in steps_android_1:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(11)

    p_c2 = doc.add_paragraph()
    p_c2.paragraph_format.space_before = Pt(4)
    p_c2.paragraph_format.space_after = Pt(4)
    r = p_c2.add_run("Cách 2 (Thủ công):")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)

    steps_android_2 = [
        "Mở link bằng Google Chrome.",
        "Bấm vào biểu tượng 3 dấu chấm (⋮) ở góc trên bên phải màn hình.",
        "Chọn \"Cài đặt ứng dụng\" hoặc \"Thêm vào màn hình chính\" (Add to Home Screen).",
        "Bấm \"Cài đặt\" để xác nhận."
    ]
    for text in steps_android_2:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(11)

    # 1.3 PC
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("1.3. Trên Máy tính (PC / Laptop Windows & Mac)")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    steps_pc = [
        ("Bước 1: ", "Mở link trên trình duyệt Google Chrome hoặc Microsoft Edge."),
        ("Bước 2: ", "Nhìn lên thanh địa chỉ trang web (thanh URL) ở góc phải, chọn biểu tượng Cài đặt ứng dụng 💻 (hoặc bấm dấu 3 chấm góc trên phải -> chọn Cài đặt Gia Phả Phạm Tộc...)."),
        ("Bước 3: ", "Nhấn \"Cài đặt\" (Install)."),
        ("🎉 Kết quả: ", "Ứng dụng sẽ tự động tạo lối tắt (shortcut) trên màn hình Desktop và mở ra dưới dạng cửa sổ độc lập mượt mà như phần mềm máy tính.")
    ]
    for b, text in steps_pc:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_b = p.add_run(b)
        r_b.bold = True
        r_b.font.name = 'Arial'
        r_b.font.size = Pt(11)
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(11)

    # --- PHẦN 2 ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("🛠️ PHẦN 2: HƯỚNG DẪN SỬ DỤNG CÁC TÍNH NĂNG CHI TIẾT")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = RGBColor(139, 26, 26)

    # 2.1 Calendar Integration
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("🔔 2.1. Tích Hợp Lịch Giỗ Vào Lịch Điện Thoại (Google Calendar, Apple Calendar)")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    p_b_desc = doc.add_paragraph()
    p_b_desc.paragraph_format.space_after = Pt(4)
    r = p_b_desc.add_run("Hệ thống cho phép bạn đồng bộ trực tiếp các ngày giỗ, sinh nhật quan trọng vào ứng dụng Lịch mặc định trên điện thoại của mình. Khi đã đồng bộ, điện thoại sẽ báo thức nhắc nhở bạn tự động (trước 7 ngày, 3 ngày, 1 ngày và đúng 08:00 sáng ngày diễn ra).")
    r.font.name = 'Arial'
    r.font.size = Pt(11)

    p_steps_title = doc.add_paragraph()
    p_steps_title.paragraph_format.space_after = Pt(4)
    r = p_steps_title.add_run("Các bước thực hiện:")
    r.bold = True
    r.font.name = 'Arial'

    steps_bell = [
        "Mở sang Tab \"Lịch giỗ\".",
        "Tìm đến ngày giỗ hoặc sinh nhật bạn muốn lưu.",
        "Bấm vào nút \"Thêm vào lịch\" (kế bên nhãn đếm ngược số ngày).",
        "Hệ thống sẽ tải xuống một tệp lịch (đuôi .ics).",
        "Mở tệp vừa tải về, điện thoại sẽ tự động mở ứng dụng Lịch (Apple Calendar trên iPhone hoặc Google Calendar trên Android).",
        "Nhấn \"Thêm\" (Add) để lưu sự kiện vào lịch điện thoại."
    ]
    for idx, text in enumerate(steps_bell, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(3)
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(11)

    p_rem = doc.add_paragraph()
    p_rem.paragraph_format.space_before = Pt(4)
    p_rem.paragraph_format.space_after = Pt(4)
    r_t2 = p_rem.add_run("Đặc biệt: Bạn có thể bấm trực tiếp vào tên người trong danh sách lịch giỗ để mở xem ngay tiểu sử chi tiết.")
    r_t2.font.name = 'Arial'
    r_t2.bold = True
    r_t2.font.color.rgb = RGBColor(0, 100, 0)

    # 2.2 Theme
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("🌓 2.2. Thay Đổi Giao Diện (Nền Tối / Nền Sáng)")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    p_theme = doc.add_paragraph()
    p_theme.paragraph_format.space_after = Pt(4)
    r = p_theme.add_run("Nhấp vào biểu tượng Mặt trời ☀️ / Mặt trăng 🌙 trên thanh menu góc trên bên phải để chuyển đổi:")
    r.font.name = 'Arial'

    themes = [
        ("🌙 Giao diện Tối (Dark Mode): ", "Phong cách sang trọng, ấm cúng với tông màu sơn son thếp vàng, dịu mắt khi xem ban đêm."),
        ("☀️ Giao diện Sáng (Light Mode): ", "Phong cách trang nhã, rõ ràng, dễ quan sát ngoài trời hoặc môi trường sáng.")
    ]
    for b, text in themes:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_b = p.add_run(b)
        r_b.bold = True
        r_b.font.name = 'Arial'
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'

    # 2.3 List
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("📋 2.3. Xem & Tra Cứu Danh Sách Thành Viên (Tab \"Danh sách\")")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    list_feats = [
        ("Tìm kiếm nhanh: ", "Nhập tên thành viên, tên thường gọi hoặc biệt danh vào ô tìm kiếm."),
        ("Bộ lọc thông minh: ", "Lọc theo Đời / Thế hệ (Đời 1, Đời 2, Đời 3...) hoặc Trạng thái (Còn sống / Đã mất)."),
        ("Xem chi tiết: ", "Bấm vào tên bất kỳ thành viên nào để mở bảng hồ sơ đầy đủ: Họ tên, tên tự, tên hiệu, ngày sinh/ngày mất (Âm & Dương lịch), mối quan hệ gia đình (cha mẹ, vợ chồng, con cái), nơi an táng và ghi chú dòng họ.")
    ]
    for b, text in list_feats:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_b = p.add_run(b)
        r_b.bold = True
        r_b.font.name = 'Arial'
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'

    # 2.4 Tree
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("🌿 2.4. Xem Sơ Đồ Phả Hệ Trực Quan (Tab \"Sơ đồ\")")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    tree_feats = [
        ("Hiển thị trực quan: ", "Cho phép hình dung rõ ràng cấu trúc gia tộc theo từng nhánh, từng đời."),
        ("Chuẩn truyền thống: ", "Sơ đồ đã được tối ưu hiển thị theo nguyên tắc truyền thống (từ Phải qua Trái): Vợ cả ➔ Vợ 2, Con trưởng ➔ Con út."),
        ("Đóng / Mở nhánh: ", "Sơ đồ ban đầu hiển thị rút gọn để dễ nhìn. Bạn có thể bấm nút mũi tên Lên/Xuống ở dưới mỗi người để Mở rộng hoặc Thu gọn các thế hệ con cháu của người đó."),
        ("Kéo rê: ", "Dùng tay vuốt (trên điện thoại) hoặc giữ chuột kéo (trên PC) để di chuyển sơ đồ."),
        ("Phóng to / Thu nhỏ: ", "Dùng 2 ngón tay chụm/xòe hoặc lăn con trỏ chuột để thu phóng."),
        ("Xem thông tin chi tiết: ", "Bấm trực tiếp vào từng ô thành viên trên cây để xem bảng tiểu sử chi tiết.")
    ]
    for b, text in tree_feats:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_b = p.add_run(b)
        r_b.bold = True
        r_b.font.name = 'Arial'
        if b in ["Chuẩn truyền thống: ", "Đóng / Mở nhánh: "]:
            r_b.font.color.rgb = RGBColor(0, 100, 0)
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'

    # 2.5 Calendar
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("📅 2.5. Tra Cứu Lịch Giỗ & Sinh Nhật (Tab \"Lịch giỗ\")")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    cal_feats = [
        ("Tự động quy đổi: ", "Hiển thị ngày giỗ Âm lịch được chuyển sang ngày Dương lịch tương ứng của năm hiện tại."),
        ("Chuyển đổi năm/tháng: ", "Bấm mũi tên < hoặc > để xem lịch các năm, chọn từng Tháng (Th.1 -> Th.12) để xem chi tiết các sự kiện."),
        ("Xem chi tiết: ", "Bấm trực tiếp vào tên thành viên trong Lịch để xem ngay thông tin tiểu sử đầy đủ.")
    ]
    for b, text in cal_feats:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_b = p.add_run(b)
        r_b.bold = True
        r_b.font.name = 'Arial'
        if b == "Xem chi tiết: ":
            r_b.font.color.rgb = RGBColor(0, 100, 0)
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'

    # 2.6 Stats
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("📊 2.6. Thống Kê Dòng Họ (Tab \"Thống kê\")")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(180, 100, 20)

    p_stat = doc.add_paragraph()
    p_stat.paragraph_format.space_after = Pt(4)
    r = p_stat.add_run("Cung cấp biểu đồ trực quan về tổng số thành viên, tỷ lệ Nam/Nữ, số lượng theo từng Đời và độ tuổi. ")
    r.font.name = 'Arial'
    
    r_t2 = p_stat.add_run("Đặc biệt: Ở dưới các danh sách mới cập nhật, bạn có thể bấm vào tên người để xem chi tiết.")
    r_t2.font.name = 'Arial'
    r_t2.bold = True
    r_t2.font.color.rgb = RGBColor(0, 100, 0)

    # --- PHẦN 3 ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("❓ PHẦN 3: CÂU HỎI THƯỜNG GẶP (FAQ)")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = RGBColor(139, 26, 26)

    faqs = [
        ("1. Tôi có cần đăng ký tài khoản hay đăng nhập không?", "Không cần. Tất cả thành viên trong dòng họ đều có thể truy cập ngay lập tức thông qua đường link."),
        ("2. Tại sao tôi không nhận được thông báo ngày giỗ?", "Hãy kiểm tra 2 điều: (1) Bạn đã Bật thông báo trong biểu tượng Cài đặt ⚙️ chưa; (2) Điện thoại/Trình duyệt của bạn có đang chặn quyền thông báo của trang web hay không."),
        ("3. Ứng dụng có bị mất dữ liệu khi tôi đổi điện thoại không?", "Không. Dữ liệu gia phả được lưu trữ tập trung trên hệ thống trực tuyến, bạn chỉ cần mở link trên thiết bị mới là xem được đầy đủ dữ liệu mới nhất.")
    ]

    for q, a in faqs:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(2)
        r_q = p_q.add_run(q)
        r_q.bold = True
        r_q.font.name = 'Arial'
        r_q.font.size = Pt(11)
        r_q.font.color.rgb = RGBColor(30, 30, 30)

        p_a = doc.add_paragraph()
        p_a.paragraph_format.space_after = Pt(4)
        r_a = p_a.add_run("👉 " + a)
        r_a.font.name = 'Arial'
        r_a.font.size = Pt(10.5)
        r_a.font.color.rgb = RGBColor(70, 70, 70)

    # Footer note
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_callout(
        doc,
        "Kính chúc toàn thể thành viên trong dòng họ Phạm Tộc luôn đoàn kết, an lạc và giữ gìn truyền thống gia phong!",
        "🤝 LỜI CHÚC: "
    )

    doc.save("/Users/mrdong/giaphaphamtoc/HUONG_DAN_SU_DUNG.docx")
    print("Successfully generated HUONG_DAN_SU_DUNG.docx")

if __name__ == "__main__":
    main()
